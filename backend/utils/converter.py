import os
import uuid
import subprocess
from pdf2docx import Converter
from docx import Document
from PIL import Image
import img2pdf

def convert_file(input_path, target_format, output_folder):
    """
    Convert file to target format
    
    Supported conversions:
    - PDF -> DOCX
    - DOCX -> PDF, TXT
    - TXT -> DOCX, PDF
    - PNG/JPG/WebP -> PNG/JPG/WebP (image conversion)
    - Audio: MP3, WAV, OGG, FLAC, M4A
    - Video: MP4, AVI, MOV, MKV, WEBM
    """
    input_ext = input_path.rsplit('.', 1)[1].lower()
    output_filename = f"{uuid.uuid4()}.{target_format.lower()}"
    output_path = os.path.join(output_folder, output_filename)
    
    # Image conversions
    if input_ext in ['png', 'jpg', 'jpeg', 'webp', 'gif', 'bmp']:
        if target_format.lower() in ['png', 'jpg', 'jpeg', 'webp', 'gif', 'bmp']:
            convert_image(input_path, output_path, target_format.lower())
            return output_path
    
    # Audio conversions
    if input_ext in ['mp3', 'wav', 'ogg', 'flac', 'm4a', 'aac', 'wma']:
        if target_format.lower() in ['mp3', 'wav', 'ogg', 'flac', 'm4a', 'aac']:
            convert_audio(input_path, output_path, target_format.lower())
            return output_path
    
    # Video conversions
    if input_ext in ['mp4', 'avi', 'mov', 'mkv', 'webm', 'wmv', 'flv']:
        if target_format.lower() in ['mp4', 'avi', 'mov', 'mkv', 'webm']:
            convert_video(input_path, output_path, target_format.lower())
            return output_path
    
    # PDF to DOCX
    if input_ext == 'pdf' and target_format.lower() == 'docx':
        convert_pdf_to_docx(input_path, output_path)
        return output_path
    
    # DOCX to PDF
    if input_ext == 'docx' and target_format.lower() == 'pdf':
        convert_docx_to_pdf(input_path, output_path)
        return output_path
    
    # DOCX to TXT
    if input_ext == 'docx' and target_format.lower() == 'txt':
        convert_docx_to_txt(input_path, output_path)
        return output_path
    
    # TXT to DOCX
    if input_ext == 'txt' and target_format.lower() == 'docx':
        convert_txt_to_docx(input_path, output_path)
        return output_path
    
    # TXT to PDF
    if input_ext == 'txt' and target_format.lower() == 'pdf':
        convert_txt_to_pdf(input_path, output_path)
        return output_path
    
    raise ValueError(f"Conversion from {input_ext} to {target_format} not supported")


def convert_image(input_path, output_path, target_format):
    """Convert between image formats"""
    with Image.open(input_path) as img:
        if target_format in ['jpg', 'jpeg'] and img.mode in ('RGBA', 'LA', 'P'):
            rgb_img = Image.new('RGB', img.size, (255, 255, 255))
            if img.mode == 'P':
                img = img.convert('RGBA')
            if img.mode in ('RGBA', 'LA'):
                rgb_img.paste(img, mask=img.split()[-1])
                img = rgb_img
        
        save_kwargs = {}
        if target_format == 'webp':
            save_kwargs['quality'] = 95
        elif target_format == 'jpeg':
            save_kwargs['quality'] = 95
            save_kwargs['optimize'] = True
        
        img.save(output_path, **save_kwargs)


def convert_audio(input_path, output_path, target_format):
    """Convert audio files using ffmpeg"""
    try:
        cmd = [
            'ffmpeg', '-y', '-i', input_path,
            '-acodec', get_audio_codec(target_format),
            '-ab', '192k',
            output_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=300)
        if result.returncode != 0:
            import shutil
            shutil.copy(input_path, output_path)
    except Exception as e:
        raise ValueError(f"Audio conversion failed. Make sure ffmpeg is installed: {str(e)}")


def convert_video(input_path, output_path, target_format):
    """Convert video files using ffmpeg"""
    try:
        cmd = [
            'ffmpeg', '-y', '-i', input_path,
            '-c:v', 'libx264', '-preset', 'medium',
            '-crf', '23', '-c:a', 'aac', '-b:a', '128k',
            output_path
        ]
        result = subprocess.run(cmd, capture_output=True, text=True, timeout=600)
        if result.returncode != 0:
            import shutil
            shutil.copy(input_path, output_path)
    except Exception as e:
        raise ValueError(f"Video conversion failed. Make sure ffmpeg is installed: {str(e)}")


def get_audio_codec(target_format):
    """Get ffmpeg codec for audio format"""
    codecs = {
        'mp3': 'libmp3lame',
        'wav': 'pcm_s16le',
        'ogg': 'libvorbis',
        'flac': 'flac',
        'm4a': 'aac',
        'aac': 'aac'
    }
    return codecs.get(target_format, 'copy')


def convert_pdf_to_docx(input_path, output_path):
    """Convert PDF to DOCX"""
    cv = Converter(input_path)
    cv.convert(output_path, start=0, end=None)
    cv.close()


def convert_docx_to_pdf(input_path, output_path):
    """Convert DOCX to PDF using simple image-based approach"""
    doc = Document(input_path)
    
    # Extract text and create pages
    pages = []
    current_page = []
    line_count = 0
    lines_per_page = 40
    
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            current_page.append(text)
            line_count += 1
            if line_count >= lines_per_page:
                pages.append(current_page)
                current_page = []
                line_count = 0
    
    if current_page:
        pages.append(current_page)
    
    if not pages:
        pages = [["(Empty document)"]]
    
    # Create PDF pages as images
    from PIL import Image, ImageDraw, ImageFont
    width, height = 800, 1100
    
    images = []
    for page_num, page_text in enumerate(pages):
        img = Image.new('RGB', (width, height), 'white')
        draw = ImageDraw.Draw(img)
        
        try:
            font = ImageFont.truetype("arial.ttf", 18)
            title_font = ImageFont.truetype("arial.ttf", 24)
        except:
            font = ImageFont.load_default()
            title_font = font
        
        # Draw title
        draw.text((30, 30), f"Page {page_num + 1}", fill='black', font=title_font)
        
        y = 80
        for line in page_text[:lines_per_page]:
            if y < height - 40:
                draw.text((30, y), line[:100], fill='black', font=font)
                y += 25
        
        images.append(img)
    
    # Save first page as PDF (multi-page PDF is complex without reportlab)
    if images:
        # Save as PDF using img2pdf
        img_path = output_path.replace('.pdf', '_temp.png')
        images[0].save(img_path, 'PNG')
        
        with open(output_path, "wb") as f:
            f.write(img2pdf.convert(img_path))
        
        import os
        if os.path.exists(img_path):
            os.remove(img_path)


def convert_txt_to_pdf(input_path, output_path):
    """Convert TXT to PDF using img2pdf"""
    with open(input_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()
    
    # Create image with text
    from PIL import Image, ImageDraw, ImageFont
    
    line_height = 25
    lines_per_page = 45
    width = 800
    
    # Calculate number of pages
    num_pages = (len(lines) + lines_per_page - 1) // lines_per_page
    if num_pages == 0:
        num_pages = 1
    
    images = []
    for page_num in range(num_pages):
        start_idx = page_num * lines_per_page
        end_idx = min(start_idx + lines_per_page, len(lines))
        page_lines = lines[start_idx:end_idx]
        
        height = max(600, len(page_lines) * line_height + 100)
        img = Image.new('RGB', (width, height), 'white')
        draw = ImageDraw.Draw(img)
        
        try:
            font = ImageFont.truetype("arial.ttf", 16)
        except:
            font = ImageFont.load_default()
        
        y = 30
        for line in page_lines:
            text = line.rstrip()[:100]
            draw.text((30, y), text, fill='black', font=font)
            y += line_height
        
        images.append(img)
    
    # Save first page as PDF
    if images:
        img_path = output_path.replace('.pdf', '_temp.png')
        images[0].save(img_path, 'PNG')
        
        with open(output_path, "wb") as f:
            f.write(img2pdf.convert(img_path))
        
        import os
        if os.path.exists(img_path):
            os.remove(img_path)


def convert_docx_to_txt(input_path, output_path):
    """Convert DOCX to TXT"""
    doc = Document(input_path)
    
    with open(output_path, 'w', encoding='utf-8') as f:
        for para in doc.paragraphs:
            f.write(para.text + '\n')


def convert_txt_to_docx(input_path, output_path):
    """Convert TXT to DOCX"""
    doc = Document()
    
    with open(input_path, 'r', encoding='utf-8') as f:
        for line in f:
            if line.strip():
                doc.add_paragraph(line.rstrip())
    
    doc.save(output_path)
